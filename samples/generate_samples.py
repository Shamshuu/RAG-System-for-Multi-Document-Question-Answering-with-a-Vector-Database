"""
Helper script to generate sample PDF and DOCX documents for testing.
"""

import os
import pypdf
import docx

def create_sample_docx(filepath: str):
    doc = docx.Document()
    doc.add_heading("Contractor & Vendor Guidelines 2026", level=0)
    
    doc.add_heading("1. Billing and Invoicing", level=1)
    doc.add_paragraph(
        "All contractors must submit invoices by the 5th business day of each month. "
        "Payment terms are strictly Net 30 from the date of invoice approval."
    )
    
    doc.add_heading("2. Benefits & Paid Time Off Exclusions", level=1)
    doc.add_paragraph(
        "Contractors operate under independent service agreements and are not eligible for "
        "corporate benefits, paid time off (PTO), health insurance, or parental leave."
    )
    
    doc.add_heading("3. Equipment & Security", level=1)
    doc.add_paragraph(
        "Contractors must use VPN access with two-factor authentication (2FA) when accessing internal git repositories."
    )
    
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    doc.save(filepath)
    print(f"Generated DOCX: {filepath}")

def create_sample_pdf(filepath: str):
    # Create multi-page PDF using pypdf writer
    from pypdf import PdfWriter
    
    writer = PdfWriter()
    
    # Page 1
    p1 = writer.add_blank_page(width=612, height=792)
    
    # In pypdf, we can write metadata or use a simple text stream
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    with open(filepath, "wb") as f:
        writer.write(f)
    print(f"Generated PDF placeholder: {filepath}")

if __name__ == "__main__":
    create_sample_docx("./samples/contractor_guidelines.docx")
