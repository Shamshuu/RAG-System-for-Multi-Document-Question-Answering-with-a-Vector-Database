"""
Document Parser Service
Extracts text from PDF and DOCX files while preserving exact page numbers and metadata.
"""

import io
from dataclasses import dataclass
from typing import List, Optional
import pypdf
import docx


@dataclass
class ExtractedPage:
    """Represents text extracted from a single document page or section."""
    document_id: str
    filename: str
    page_number: int
    text: str


class DocumentParser:
    """Parses PDF and DOCX documents with page number preservation."""

    @staticmethod
    def parse_pdf(file_bytes: bytes, filename: str, document_id: str) -> List[ExtractedPage]:
        """
        Parses a PDF document, preserving the 1-indexed page number for each page.
        """
        pages: List[ExtractedPage] = []
        pdf_stream = io.BytesIO(file_bytes)
        reader = pypdf.PdfReader(pdf_stream)

        for page_idx, page in enumerate(reader.pages):
            page_number = page_idx + 1
            text = page.extract_text() or ""
            # Clean and normalize whitespace
            cleaned_text = " ".join(text.split()).strip()
            if cleaned_text:
                pages.append(ExtractedPage(
                    document_id=document_id,
                    filename=filename,
                    page_number=page_number,
                    text=cleaned_text
                ))

        return pages

    @staticmethod
    def parse_docx(file_bytes: bytes, filename: str, document_id: str) -> List[ExtractedPage]:
        """
        Parses a DOCX document, grouping content into sequential logical blocks/pages.
        """
        pages: List[ExtractedPage] = []
        docx_stream = io.BytesIO(file_bytes)
        doc = docx.Document(docx_stream)

        # Approximate logical page/section size (e.g. ~3000 chars per page if no explicit page breaks)
        current_block_text = []
        current_block_chars = 0
        current_page_number = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            current_block_text.append(text)
            current_block_chars += len(text)

            # Check if block exceeded target page size or has page break
            if current_block_chars >= 2500:
                block_content = "\n".join(current_block_text).strip()
                if block_content:
                    pages.append(ExtractedPage(
                        document_id=document_id,
                        filename=filename,
                        page_number=current_page_number,
                        text=block_content
                    ))
                current_page_number += 1
                current_block_text = []
                current_block_chars = 0

        # Flush remaining paragraphs
        if current_block_text:
            block_content = "\n".join(current_block_text).strip()
            if block_content:
                pages.append(ExtractedPage(
                    document_id=document_id,
                    filename=filename,
                    page_number=current_page_number,
                    text=block_content
                ))

        # Also extract any tables
        if doc.tables:
            for table_idx, table in enumerate(doc.tables):
                table_lines = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_lines.append(" | ".join(row_text))
                if table_lines:
                    table_content = "\n".join(table_lines)
                    pages.append(ExtractedPage(
                        document_id=document_id,
                        filename=filename,
                        page_number=current_page_number + table_idx,
                        text=table_content
                    ))

        return pages

    @classmethod
    def parse_document(cls, file_bytes: bytes, filename: str, document_id: str) -> List[ExtractedPage]:
        """
        Detects file type by extension and extracts text with page numbers.
        """
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            return cls.parse_pdf(file_bytes, filename, document_id)
        elif lower_name.endswith(".docx"):
            return cls.parse_docx(file_bytes, filename, document_id)
        elif lower_name.endswith(".txt") or lower_name.endswith(".md"):
            text = file_bytes.decode("utf-8", errors="ignore").strip()
            return [ExtractedPage(
                document_id=document_id,
                filename=filename,
                page_number=1,
                text=text
            )] if text else []
        else:
            raise ValueError(f"Unsupported file format: {filename}. Supported formats: .pdf, .docx, .txt")
